"""
Convert all Annex27 policy markdown files to branded Word (.docx) documents.
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"C:\Users\raoul\Documents\annex27\beleidspakket"
OUT = r"C:\Users\raoul\Documents\annex27\beleidspakket\word"

TEAL = RGBColor(0x0D, 0x94, 0x88)
NAVY = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)
LIGHT_GRAY = RGBColor(0x94, 0xA3, 0xB8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def create_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = GRAY

    # Heading styles
    for i, (size, color) in enumerate([(16, NAVY), (13, NAVY), (11, TEAL)], 1):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        if i == 1:
            h.paragraph_format.space_after = Pt(6)

    return doc

def add_header(doc, title):
    # Logo line
    p = doc.add_paragraph()
    run = p.add_run('annex')
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.name = 'Calibri'
    run2 = p.add_run('27')
    run2.font.size = Pt(14)
    run2.font.color.rgb = TEAL
    run2.font.bold = True
    run2.font.name = 'Calibri'

    run3 = p.add_run('    ISO 27001:2022 Beleidspakket')
    run3.font.size = Pt(8)
    run3.font.color.rgb = LIGHT_GRAY
    run3.font.name = 'Calibri'

    # Teal line
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(16)
    # Add a border line via paragraph border
    pBdr = p2._element.get_or_add_pPr()
    bdr = pBdr.makeelement(qn('w:pBdr'), {})
    bottom = bdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '1',
        qn('w:color'): '0D9488'
    })
    bdr.append(bottom)
    pBdr.append(bdr)

def parse_md_line(line):
    """Simple markdown to text parsing."""
    line = line.strip()
    # Bold
    line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
    # Code
    line = re.sub(r'`(.+?)`', r'\1', line)
    # Links
    line = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', line)
    return line

def md_to_docx(md_content, output_path, filename):
    doc = create_doc()
    lines = md_content.split('\n')

    in_table = False
    table_rows = []
    in_code = False
    first_h1 = True

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            in_code = not in_code
            if not in_code and table_rows:
                pass
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(stripped)
            p.style.font.name = 'Consolas'
            p.style.font.size = Pt(9)
            i += 1
            continue

        # Empty lines
        if not stripped:
            if in_table and table_rows:
                # Flush table
                create_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Table rows
        if '|' in stripped and stripped.startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Skip separator rows
            if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table and table_rows:
                create_table(doc, table_rows)
                table_rows = []
                in_table = False

        # Headings
        if stripped.startswith('# '):
            text = parse_md_line(stripped[2:])
            if first_h1:
                add_header(doc, text)
                first_h1 = False
            doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith('## '):
            text = parse_md_line(stripped[3:])
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith('### '):
            text = parse_md_line(stripped[4:])
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # List items
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = parse_md_line(stripped[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            p.style.font.size = Pt(10)
            i += 1
            continue

        if re.match(r'^\d+\.', stripped):
            text = parse_md_line(re.sub(r'^\d+\.\s*', '', stripped))
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        text = parse_md_line(stripped)
        if text:
            p = doc.add_paragraph(text)
            # Check for bold parts in original
            if '**' in lines[i]:
                pass  # Already parsed out

        i += 1

    # Flush remaining table
    if table_rows:
        create_table(doc, table_rows)

    # Footer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('© 2026 Annex27 — annex27.nl — Alle rechten voorbehouden')
    run.font.size = Pt(7)
    run.font.color.rgb = LIGHT_GRAY

    doc.save(output_path)

def create_table(doc, rows):
    if not rows or not rows[0]:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(parse_md_line(cell_text))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'

            if r_idx == 0:
                # Header row
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(8.5)
                # Dark background
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): '0F172A',
                    qn('w:val'): 'clear'
                })
                shading.append(shd)
            else:
                run.font.color.rgb = GRAY
                if r_idx % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'F8FAFC',
                        qn('w:val'): 'clear'
                    })
                    shading.append(shd)

def process_dir(subdir):
    src = os.path.join(BASE, subdir)
    dst = os.path.join(OUT, subdir)
    os.makedirs(dst, exist_ok=True)

    count = 0
    for f in sorted(os.listdir(src)):
        if not f.endswith('.md'):
            continue
        with open(os.path.join(src, f), 'r', encoding='utf-8') as fh:
            content = fh.read()

        docx_path = os.path.join(dst, f.replace('.md', '.docx'))
        md_to_docx(content, docx_path, f.replace('.md', ''))
        count += 1
        print(f"  {f} -> .docx")

    return count

print("Converting to Word (.docx)...\n")

total = 0
print("=== BASISPAKKET ===")
total += process_dir("basispakket")

print("\n=== PREMIUM ===")
total += process_dir("premium")

print("\n=== WERKINSTRUCTIES ===")
total += process_dir("werkinstructies")

print(f"\nDone! {total} Word documents created in: {OUT}")
